import requests
from bs4 import BeautifulSoup
from docx import Document
import os

# Step 1: Define the URL of the website you want to scrape
url = 'http://books.toscrape.com/catalogue/page-1.html'

# Step 2: Send a GET request to fetch the webpage content
response = requests.get(url)

# Step 3: Check if the request was successful
if response.status_code == 200:
    # Step 4: Parse the HTML content using BeautifulSoup
    soup = BeautifulSoup(response.text, 'html.parser')
    
    # Step 5: Create a new Document
    doc = Document()
    doc.add_heading('Book Titles and Ratings', level=1)

    # Step 6: Find all book containers on the page
    books = soup.find_all('article', class_='product_pod')

    # Step 7: Loop through each book and extract title and rating
    for book in books:
        title = book.h3.a['title']  # Extract the title
        rating = book.find('p', class_='star-rating')['class'][1]  # Extract the rating class
        
        # Add title and rating to the document
        doc.add_paragraph(f'Title: {title}, Rating: {rating}')

    # Step 8: Define the path where you want to save the document
    save_directory = './Python Projects + ChatGPT/Python Projects/Web Scraper'  # Change this to your actual path
    file_name = 'book_titles_and_ratings.docx'
    
    # Save the document in the specified folder
    doc.save(os.path.join(save_directory, file_name))
    print(f"Data saved to '{os.path.join(save_directory, file_name)}'.")
else:
    print("Failed to retrieve the webpage.")